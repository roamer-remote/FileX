# Copyright (c) 2026 徐泽宇
"""docx 业务逻辑模块。

Authors:
    徐泽宇
"""

from __future__ import annotations

try:
    from docx import Document
except ModuleNotFoundError:
    Document = None

from services.extract.base import ExtractResult


def extract_docx(path: str) -> ExtractResult:
    if Document is None:
        raise ImportError("python-docx is required to extract docx files")
    doc = Document(path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return ExtractResult(text="\n\n".join(parts).strip(), engine="python-docx")
